# Resume: extract text from PDF, Word, images (AI vision), or plain text
from io import BytesIO
from pathlib import Path

import PyPDF2
import pdfplumber

from app.services.ai_service import extract_text_with_vision


def extract_text_from_docx_bytes(content: bytes) -> str:
    """Extract text from Word .docx bytes using python-docx."""
    if not content:
        return ""
    try:
        from docx import Document
        doc = Document(BytesIO(content))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """
    Extract raw text from PDF bytes. Tries pdfplumber first (better for most PDFs),
    then PyPDF2 as fallback.
    """
    if not pdf_bytes:
        return ""
    stream = BytesIO(pdf_bytes)

    # 1) Try pdfplumber (better layout/table handling, often more accurate text)
    try:
        with pdfplumber.open(stream) as pdf:
            parts = []
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    parts.append(t.strip())
            if parts:
                return "\n".join(parts)
    except Exception:
        pass

    # 2) Fallback: PyPDF2
    try:
        stream.seek(0)
        reader = PyPDF2.PdfReader(stream)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)
    except Exception:
        return ""


def extract_text_from_upload(content: bytes, filename: str, content_type: str) -> str:
    """
    Extract text from uploaded file. Supports:
    - PDF: pdfplumber + PyPDF2
    - Word: .docx via python-docx
    - Images (PNG, JPEG, etc.): OpenAI Vision (AI model)
    - Plain text: .txt, .md
    """
    if not content:
        return ""
    suffix = (Path(filename).suffix or "").lower()
    ct = (content_type or "").lower()

    # PDF
    if "pdf" in ct or suffix == ".pdf":
        return extract_text_from_pdf_bytes(content)

    # Word
    if "wordprocessingml" in ct or "vnd.openxmlformats" in ct or suffix == ".docx":
        return extract_text_from_docx_bytes(content)

    # Images → AI Vision
    if suffix in (".png", ".jpg", ".jpeg", ".gif", ".webp") or "image/" in ct:
        mime = ct if ct.startswith("image/") else f"image/{suffix.lstrip('.')}" if suffix else "image/jpeg"
        if mime == "image/jpg":
            mime = "image/jpeg"
        return extract_text_with_vision(content, mime)

    # Plain text
    if "text" in ct or suffix in (".txt", ".md"):
        return content.decode("utf-8", errors="ignore")

    return ""
